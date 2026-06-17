from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from database import get_db
from models import Meeting, MeetingAgendaItem, PlatformSave
from routers.doc_content import item_sections
from datetime import datetime, timezone, timedelta
import os

from docx import Document

router = APIRouter(prefix="/api/meetings", tags=["docx"])

DOCX_DIR = "docx"
KST = timezone(timedelta(hours=9))  # 한국 시간


@router.post("/{meeting_id}/save-docx")
def save_docx(meeting_id: int, db: Session = Depends(get_db)):
    """회의록을 Word(.docx) 파일로 저장"""
    meeting = db.query(Meeting).filter(Meeting.meeting_id == meeting_id).first()
    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")

    items = db.query(MeetingAgendaItem).filter(
        MeetingAgendaItem.meeting_id == meeting_id
    ).order_by(MeetingAgendaItem.order).all()

    if not items:
        raise HTTPException(status_code=400, detail="No agenda items found")

    os.makedirs(DOCX_DIR, exist_ok=True)
    filename = f"meeting_{meeting_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = f"{DOCX_DIR}/{filename}"

    # 회의 생성 시각(UTC 저장) → 한국 시간 변환
    created = meeting.created_at.replace(tzinfo=timezone.utc).astimezone(KST)
    date_str = created.strftime("%Y-%m-%d %H:%M")

    doc = Document()

    # 제목 + 날짜 + 참석자
    doc.add_heading(meeting.title, level=0)
    doc.add_paragraph(f"📅 {date_str}")
    if meeting.participants:
        p = doc.add_paragraph()
        p.add_run("👥 참석자: ").bold = True
        p.add_run(meeting.participants)

    # 안건 (내용·주요 의견·결정·한 일·할 일 모두 출력)
    for item in items:
        doc.add_heading(f"{item.order}. {item.agenda}", level=1)
        for label, body in item_sections(item):
            if isinstance(body, list):
                p = doc.add_paragraph()
                p.add_run(f"{label}:").bold = True
                for b in body:
                    doc.add_paragraph(b, style="List Bullet")
            else:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(body))

    doc.save(file_path)

    db.add(PlatformSave(
        meeting_id=meeting_id,
        platform="docx",
        save_status="success",
        platform_doc_id=file_path,
    ))
    db.commit()

    return {"meeting_id": meeting_id, "status": "success", "file_path": file_path}


@router.get("/{meeting_id}/download-docx")
def download_docx(meeting_id: int, db: Session = Depends(get_db)):
    """DOCX 다운로드 (없으면 자동 생성)"""
    meeting = db.query(Meeting).filter(Meeting.meeting_id == meeting_id).first()
    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")

    items = db.query(MeetingAgendaItem).filter(
        MeetingAgendaItem.meeting_id == meeting_id
    ).order_by(MeetingAgendaItem.order).all()

    if not items:
        raise HTTPException(status_code=400, detail="No agenda items found")

    os.makedirs(DOCX_DIR, exist_ok=True)
    filename = f"meeting_{meeting_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = f"{DOCX_DIR}/{filename}"

    # 회의 생성 시각(UTC 저장) → 한국 시간 변환
    created = meeting.created_at.replace(tzinfo=timezone.utc).astimezone(KST)
    date_str = created.strftime("%Y-%m-%d %H:%M")

    doc = Document()

    # 제목 + 날짜 + 참석자
    doc.add_heading(meeting.title, level=0)
    doc.add_paragraph(f"📅 {date_str}")
    if meeting.participants:
        p = doc.add_paragraph()
        p.add_run("👥 참석자: ").bold = True
        p.add_run(meeting.participants)

    # 안건 (내용·주요 의견·결정·한 일·할 일 모두 출력)
    for item in items:
        doc.add_heading(f"{item.order}. {item.agenda}", level=1)
        for label, body in item_sections(item):
            if isinstance(body, list):
                p = doc.add_paragraph()
                p.add_run(f"{label}:").bold = True
                for b in body:
                    doc.add_paragraph(b, style="List Bullet")
            else:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(body))

    doc.save(file_path)

    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )