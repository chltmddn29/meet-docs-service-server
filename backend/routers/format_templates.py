from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
from database import get_db
from models import FormatTemplate, Transcript, Meeting, MeetingAgendaItem
from routers.groq_client import client, TEXT_MODEL
from routers.doc_content import item_sections
from pydantic import BaseModel
from io import BytesIO
from datetime import timezone, timedelta
import os
import re
import json

KST = timezone(timedelta(hours=9))
# {{...}} 플레이스홀더 패턴
_PLACEHOLDER_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")

router = APIRouter(prefix="/api/format-templates", tags=["format-templates"])


def _extract_hwpx(data: bytes) -> str:
    """HWPX(최신 한글, ZIP+XML)에서 본문 텍스트 추출. 표준 라이브러리만 사용."""
    import zipfile
    import re
    import html

    parts = []
    with zipfile.ZipFile(BytesIO(data)) as z:
        # 본문은 Contents/section0.xml, section1.xml ...
        names = sorted(
            n for n in z.namelist()
            if n.lower().startswith("contents/") and n.lower().endswith(".xml")
        )
        for n in names:
            xml = z.read(n).decode("utf-8", errors="ignore")
            # 텍스트는 <hp:t> ... </hp:t> 안에 들어있음
            for m in re.findall(r"<hp:t[^>]*>(.*?)</hp:t>", xml, re.DOTALL):
                clean = re.sub(r"<[^>]+>", "", m)        # 내부 태그 제거
                parts.append(html.unescape(clean))
    return "\n".join(p for p in parts if p.strip())


def _extract_hwp(data: bytes) -> str:
    """HWP5(구형 한글, OLE 바이너리)에서 본문 텍스트 추출. olefile 필요."""
    import zlib
    import struct
    try:
        import olefile
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="HWP 처리 모듈이 준비되지 않았어요 (서버에 olefile 설치 필요).",
        )

    try:
        ole = olefile.OleFileIO(BytesIO(data))
    except Exception:
        raise HTTPException(status_code=400, detail="HWP 파일을 열 수 없습니다 (손상되었거나 형식 오류).")

    try:
        dirs = ole.listdir()
        if ["FileHeader"] not in dirs:
            raise HTTPException(status_code=400, detail="올바른 HWP 파일이 아닙니다.")

        header = ole.openstream("FileHeader").read()
        is_compressed = bool(header[36] & 1)

        # BodyText/Section0, Section1 ... 수집
        section_nums = sorted(
            int(d[1][len("Section"):])
            for d in dirs
            if len(d) > 1 and d[0] == "BodyText" and d[1].startswith("Section")
        )

        out = []
        for num in section_nums:
            stream = ole.openstream(f"BodyText/Section{num}").read()
            if is_compressed:
                stream = zlib.decompress(stream, -15)
            out.append(_parse_hwp_section(stream, struct))
        return "\n".join(out)
    finally:
        ole.close()


def _parse_hwp_section(buf: bytes, struct) -> str:
    """HWP 섹션 레코드를 순회하며 문단 텍스트(PARA_TEXT, tag=67)만 추출."""
    HWPTAG_PARA_TEXT = 67
    parts = []
    i, size = 0, len(buf)
    while i + 4 <= size:
        header = struct.unpack_from("<I", buf, i)[0]
        rec_type = header & 0x3FF
        rec_len = (header >> 20) & 0xFFF
        data_start = i + 4
        # size가 0xFFF면 다음 4바이트가 실제 길이
        if rec_len == 0xFFF:
            rec_len = struct.unpack_from("<I", buf, i + 4)[0]
            data_start = i + 8
        if rec_type == HWPTAG_PARA_TEXT:
            raw = buf[data_start:data_start + rec_len]
            s = raw.decode("utf-16-le", errors="ignore")
            # 인라인 제어문자 제거(개행·탭은 유지)
            s = "".join(c for c in s if c in "\n\t" or ord(c) >= 32)
            if s.strip():
                parts.append(s)
        i = data_start + rec_len
    return "\n".join(parts)


def _extract_text(filename: str, data: bytes) -> str:
    """업로드 파일에서 본문 텍스트 추출 (docx/hwp/hwpx/md/txt 지원)"""
    ext = os.path.splitext(filename)[1].lower()
    if ext in (".txt", ".md"):
        return data.decode("utf-8", errors="ignore")
    if ext == ".docx":
        from docx import Document
        doc = Document(BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        # 표 안의 텍스트도 포함
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    if ext == ".hwpx":
        return _extract_hwpx(data)
    if ext == ".hwp":
        return _extract_hwp(data)
    raise HTTPException(
        status_code=400,
        detail=f"지원하지 않는 형식입니다: {ext or '확장자 없음'} (docx/hwp/hwpx/md/txt 가능)",
    )


def _serialize(t: FormatTemplate, preview: bool = False) -> dict:
    content = t.content or ""
    return {
        "format_template_id": t.format_template_id,
        "name": t.name,
        "content": content[:200] if preview else content,
        "source_filename": t.source_filename,
        "created_at": t.created_at.isoformat() if t.created_at else None,
    }


# 1. 서식 템플릿 목록 (본문은 미리보기 200자)
@router.get("")
def list_format_templates(db: Session = Depends(get_db)):
    rows = (
        db.query(FormatTemplate)
        .order_by(FormatTemplate.created_at.desc())
        .all()
    )
    return [_serialize(t, preview=True) for t in rows]


# 2. 파일 업로드 → 텍스트 추출 → 서식 템플릿 생성
@router.post("/upload")
async def upload_format_template(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    try:
        data = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"파일 읽기 실패: {e}")
    if not data:
        raise HTTPException(status_code=400, detail="빈 파일입니다")

    text = _extract_text(file.filename or "", data)
    if not text.strip():
        raise HTTPException(
            status_code=400,
            detail="파일에서 텍스트를 추출하지 못했습니다 (빈 문서이거나 형식 문제).",
        )

    name = os.path.splitext(file.filename or "서식")[0]
    t = FormatTemplate(name=name, content=text, source_filename=file.filename)
    try:
        db.add(t)
        db.commit()
        db.refresh(t)
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"서식 저장 실패: {e}")
    return _serialize(t, preview=True)


# 2-1. 예시 서식 추가 (플레이스홀더 데모용) — 이미 있으면 건너뜀
_EXAMPLE_TEMPLATES = {
    "표준 회의록 (예시)": """# {{제목}}

- 일시: {{날짜}}
- 참석자: {{참석자}}

## 안건
{{안건}}

## 논의 내용
{{내용}}

## 결정 사항
{{결정}}

## 할 일
{{할일}}
""",
    "실무 액션 중심 (예시)": """# {{제목}} 회의 결과

📅 {{날짜}}  |  👥 {{참석자}}

## ✅ 한 일
{{한일}}

## 📋 할 일
{{할일}}

## 💡 결정 사항
{{결정}}
""",
    "의견·발언 정리형 (예시)": """# {{제목}}

일시: {{날짜}}
참석: {{참석자}}

## 주요 의견
{{의견}}

## 발언자별 정리
{{발언자}}

## 결정
{{결정}}

## 액션 아이템
{{할일}}
""",
}


@router.post("/examples")
def add_example_templates(db: Session = Depends(get_db)):
    """플레이스홀더가 든 예시 서식들을 추가. 같은 이름이 이미 있으면 건너뜀."""
    existing = {t.name for t in db.query(FormatTemplate.name).all()}
    created = 0
    try:
        for name, content in _EXAMPLE_TEMPLATES.items():
            if name in existing:
                continue
            db.add(FormatTemplate(
                name=name, content=content, source_filename="example"
            ))
            created += 1
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"예시 추가 실패: {e}")
    return {"added": created, "total_examples": len(_EXAMPLE_TEMPLATES)}


# 3. 서식 템플릿 삭제
@router.delete("/{format_template_id}")
def delete_format_template(format_template_id: int, db: Session = Depends(get_db)):
    t = (
        db.query(FormatTemplate)
        .filter(FormatTemplate.format_template_id == format_template_id)
        .first()
    )
    if not t:
        raise HTTPException(status_code=404, detail="Format template not found")
    db.delete(t)
    db.commit()
    return {"message": "Format template deleted"}


class GenerateRequest(BaseModel):
    meeting_id: int
    format_template_id: int


def _jload(v) -> list:
    if not v:
        return []
    try:
        x = json.loads(v)
        return x if isinstance(x, list) else []
    except (json.JSONDecodeError, TypeError):
        return []


def _minutes_body(items) -> str:
    """안건별 정리를 읽기 좋은 본문 텍스트로."""
    parts = []
    for it in items:
        parts.append(f"## {it.order}. {it.agenda}")
        for label, body in item_sections(it):
            if isinstance(body, list):
                parts.append(f"**{label}**")
                parts.extend(f"- {b}" for b in body)
            else:
                parts.append(f"**{label}**\n{body}")
    return "\n".join(parts)


def _meeting_values(meeting, items, raw_text):
    """플레이스홀더 → 실제 값 매핑. (정확 치환용)"""
    date_str = ""
    if meeting.created_at:
        date_str = (
            meeting.created_at.replace(tzinfo=timezone.utc)
            .astimezone(KST)
            .strftime("%Y-%m-%d %H:%M")
        )

    decisions, actions, completed, discussions, speakers = [], [], [], [], []
    for it in items:
        if it.decision:
            decisions.append(it.decision)
        actions += _jload(it.action_items)
        completed += _jload(getattr(it, "completed_items", None))
        discussions += _jload(getattr(it, "discussions", None))
        speakers += _jload(getattr(it, "speaker_points", None))

    def bullets(xs):
        return "\n".join(f"- {x}" for x in xs)

    body = _minutes_body(items)
    values = {
        "title": meeting.title or "",
        "date": date_str,
        "participants": meeting.participants or "",
        "agenda": "\n".join(f"{it.order}. {it.agenda}" for it in items),
        "content": body,
        "body": body,
        "minutes": body,
        "decisions": bullets(decisions),
        "action_items": bullets(actions),
        "todos": bullets(actions),
        "completed": bullets(completed),
        "discussions": bullets(discussions),
        "speakers": bullets(speakers),
        "raw_text": raw_text or "",
    }
    # 한글 별칭 → 표준 키
    aliases = {
        "제목": "title", "회의제목": "title", "회의명": "title",
        "날짜": "date", "일시": "date", "회의일시": "date",
        "참석자": "participants", "참가자": "participants",
        "안건": "agenda", "안건목록": "agenda",
        "내용": "content", "회의록": "content", "본문": "content", "회의내용": "content",
        "결정": "decisions", "결정사항": "decisions",
        "할일": "action_items", "할일목록": "action_items", "액션아이템": "action_items",
        "한일": "completed", "완료": "completed", "완료사항": "completed",
        "의견": "discussions", "주요의견": "discussions",
        "발언자": "speakers", "발언자별정리": "speakers",
        "원본": "raw_text", "전문": "raw_text", "원본텍스트": "raw_text",
    }
    return values, aliases


def _ai_fill_unknown(tokens, meeting, raw_text) -> dict:
    """표준 키에 없는 사용자 정의 플레이스홀더를 회의 내용 기반으로 AI가 채움."""
    listing = "\n".join(f"- {t}" for t in tokens)
    try:
        resp = client.chat.completions.create(
            model=TEXT_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "회의 내용을 바탕으로 요청된 각 항목의 값을 채우는 AI입니다. "
                        "반드시 {\"항목명\": \"값\"} 형태의 순수 JSON 객체만 응답하세요. "
                        "회의에 없는 내용은 지어내지 말고 빈 문자열로 두세요."
                    ),
                },
                {
                    "role": "user",
                    "content": (
                        f"[회의 제목]\n{meeting.title}\n\n"
                        f"[회의 내용]\n{raw_text}\n\n"
                        f"[채울 항목들]\n{listing}\n\n"
                        "각 항목에 들어갈 값을 JSON으로 주세요."
                    ),
                },
            ],
            temperature=0.3,
        )
        text = resp.choices[0].message.content.strip()
        if "```" in text:
            text = text.split("```")[1]
            if text.lstrip().lower().startswith("json"):
                text = text.lstrip()[4:]
            text = text.strip()
        data = json.loads(text)
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}  # AI 실패 시 빈 값으로 (전체 실패 방지)


def _fill_placeholders(template_text, meeting, items, raw_text) -> str:
    """서식의 {{...}} 플레이스홀더를 실제 값으로 정확히 치환."""
    values, aliases = _meeting_values(meeting, items, raw_text)
    norm_alias = {k.replace(" ", ""): v for k, v in aliases.items()}

    tokens = list({m.strip() for m in _PLACEHOLDER_RE.findall(template_text)})
    resolved, unknown = {}, []
    for tok in tokens:
        key = tok.lower().replace(" ", "").replace("_", "")
        if tok.lower() in values:
            resolved[tok] = values[tok.lower()]
        elif key in {k.replace("_", "") for k in values}:
            # 표준 키(언더스코어 무시) 매칭
            for vk in values:
                if vk.replace("_", "") == key:
                    resolved[tok] = values[vk]
                    break
        elif tok in aliases:
            resolved[tok] = values[aliases[tok]]
        elif tok.replace(" ", "") in norm_alias:
            resolved[tok] = values[norm_alias[tok.replace(" ", "")]]
        else:
            unknown.append(tok)

    if unknown:
        ai_vals = _ai_fill_unknown(unknown, meeting, raw_text)
        for tok in unknown:
            resolved[tok] = str(ai_vals.get(tok, ""))

    return _PLACEHOLDER_RE.sub(
        lambda m: str(resolved.get(m.group(1).strip(), "")), template_text
    )


# 4. 회의 원본을 서식 템플릿 형식대로 AI로 생성
@router.post("/generate")
def generate_formatted(req: GenerateRequest, db: Session = Depends(get_db)):
    """회의 transcript를 서식 템플릿 형식에 맞춰 AI 회의록으로 생성 (저장하지 않고 반환)"""
    template = (
        db.query(FormatTemplate)
        .filter(FormatTemplate.format_template_id == req.format_template_id)
        .first()
    )
    if not template:
        raise HTTPException(status_code=404, detail="Format template not found")

    meeting = (
        db.query(Meeting).filter(Meeting.meeting_id == req.meeting_id).first()
    )
    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")

    transcript = (
        db.query(Transcript)
        .filter(Transcript.meeting_id == req.meeting_id)
        .first()
    )
    raw_text = transcript.raw_text if transcript else None
    if not raw_text:
        raise HTTPException(
            status_code=400, detail="회의 원본 텍스트가 없습니다 (먼저 음성 처리 필요)"
        )

    # 서식에 {{...}} 플레이스홀더가 있으면 → 정확히 값 채우기 모드
    if _PLACEHOLDER_RE.search(template.content or ""):
        items = (
            db.query(MeetingAgendaItem)
            .filter(MeetingAgendaItem.meeting_id == req.meeting_id)
            .order_by(MeetingAgendaItem.order)
            .all()
        )
        try:
            filled = _fill_placeholders(template.content, meeting, items, raw_text)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"서식 채우기 실패: {e}")
        return {
            "meeting_id": req.meeting_id,
            "format_template_id": req.format_template_id,
            "formatted": filled,
            "mode": "placeholder",
        }

    # 플레이스홀더가 없으면 → 형식 모방 모드(기존)
    try:
        response = client.chat.completions.create(
            model=TEXT_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "당신은 회의록 작성 AI입니다. 사용자가 제공한 '예시 서식' 문서의 "
                        "구조·제목 체계·섹션 구성·말투를 그대로 본떠서, 이번 회의 내용을 같은 형식의 "
                        "회의록으로 작성하세요. 결과는 마크다운으로만 출력하고 설명이나 코드블록 표시(```)는 넣지 마세요."
                    ),
                },
                {
                    "role": "user",
                    "content": f"""[예시 서식 — 이 형식과 구조를 그대로 따라 작성]
{template.content}

[이번 회의 제목]
{meeting.title}

[이번 회의 원본 내용(STT)]
{raw_text}

위 '예시 서식'과 동일한 구조·스타일로, 이번 회의 내용을 채워 완성된 회의록을 마크다운으로 작성해주세요.""",
                },
            ],
            temperature=0.4,
        )

        content = response.choices[0].message.content.strip()
        # 혹시 코드블록으로 감싸오면 제거
        if content.startswith("```"):
            lines = content.split("\n")
            if lines and lines[0].startswith("```"):
                lines = lines[1:]
            if lines and lines[-1].strip() == "```":
                lines = lines[:-1]
            content = "\n".join(lines).strip()

        return {
            "meeting_id": req.meeting_id,
            "format_template_id": req.format_template_id,
            "formatted": content,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))
